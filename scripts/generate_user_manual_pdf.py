#!/usr/bin/env python3
"""
Generate the branded PDF of the Jana Financed Emissions Dashboard User's Manual.

Content is read from a python-docx source. Layout, cover, header/footer, and
callout/table styling are rendered via ReportLab to match the v0.1 branded look.

Usage:
    python3 scripts/generate_user_manual_pdf.py \
        --version 0.2 \
        --date "November 2026" \
        --docx docs/Jana_Financed_Emissions_Dashboard_User_Manual_Demo_v0.2_DRAFT.docx \
        --out docs/Jana_Financed_Emissions_Dashboard_User_Manual_Demo_v0.2.pdf

For future revisions (v0.3, v0.4, ...) update the docx source, then re-run
with the new --version and --date flags.
"""
from __future__ import annotations

import argparse
import os
import sys
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)

# --- Brand palette (matches v0.1 rendered PDF; see docs/STYLE_NOTES.md) -----
BRAND_GREEN       = colors.HexColor("#0F5132")   # deep forest, primary
BRAND_ACCENT      = colors.HexColor("#1B6B3D")   # slightly lighter accent
HEADER_TEXT       = colors.HexColor("#334155")   # slate grey for sub-heads
BODY_TEXT         = colors.HexColor("#1F2937")   # near-black body
RULE_GREY         = colors.HexColor("#94A3B8")   # table borders
CELL_HEAD_FILL    = colors.HexColor("#E8F1EC")   # very light green header row
NOTICE_BG         = colors.HexColor("#FEF3C7")   # amber-tinted callout bg
NOTICE_BORDER     = colors.HexColor("#F59E0B")   # amber border
FOOTER_GREY       = colors.HexColor("#64748B")

# --- Layout ------------------------------------------------------------------
PAGE_W, PAGE_H    = LETTER
MARGIN_L          = 0.85 * inch
MARGIN_R          = 0.85 * inch
MARGIN_T          = 0.85 * inch
MARGIN_B          = 0.85 * inch
CONTENT_W         = PAGE_W - MARGIN_L - MARGIN_R

FONT              = "Helvetica"
FONT_B            = "Helvetica-Bold"
FONT_I            = "Helvetica-Oblique"


# ---------------------------------------------------------------------------
# Paragraph styles
# ---------------------------------------------------------------------------
def build_styles() -> dict:
    ss = getSampleStyleSheet()
    styles = {}

    styles["Body"] = ParagraphStyle(
        "Body",
        parent=ss["BodyText"],
        fontName=FONT,
        fontSize=10.5,
        leading=14.5,
        textColor=BODY_TEXT,
        spaceAfter=6,
    )
    styles["Bullet"] = ParagraphStyle(
        "Bullet",
        parent=styles["Body"],
        leftIndent=18,
        bulletIndent=6,
        spaceAfter=3,
    )
    styles["H1"] = ParagraphStyle(
        "H1",
        parent=ss["Heading1"],
        fontName=FONT_B,
        fontSize=17,
        leading=21,
        textColor=BRAND_GREEN,
        spaceBefore=6,
        spaceAfter=10,
    )
    styles["H2"] = ParagraphStyle(
        "H2",
        parent=ss["Heading2"],
        fontName=FONT_B,
        fontSize=12.5,
        leading=15,
        textColor=BRAND_ACCENT,
        spaceBefore=10,
        spaceAfter=4,
    )
    styles["H3"] = ParagraphStyle(
        "H3",
        parent=ss["Heading3"],
        fontName=FONT_B,
        fontSize=11,
        leading=13.5,
        textColor=HEADER_TEXT,
        spaceBefore=8,
        spaceAfter=2,
    )
    styles["CoverTitle"] = ParagraphStyle(
        "CoverTitle",
        parent=ss["Title"],
        fontName=FONT_B,
        fontSize=30,
        leading=36,
        textColor=BRAND_GREEN,
        alignment=TA_CENTER,
    )
    styles["CoverSub"] = ParagraphStyle(
        "CoverSub",
        parent=ss["Title"],
        fontName=FONT_B,
        fontSize=20,
        leading=24,
        textColor=HEADER_TEXT,
        alignment=TA_CENTER,
    )
    styles["CoverTag"] = ParagraphStyle(
        "CoverTag",
        parent=ss["Normal"],
        fontName=FONT_I,
        fontSize=13,
        leading=16,
        textColor=HEADER_TEXT,
        alignment=TA_CENTER,
    )
    styles["CoverMeta"] = ParagraphStyle(
        "CoverMeta",
        parent=ss["Normal"],
        fontName=FONT,
        fontSize=12,
        leading=15,
        textColor=HEADER_TEXT,
        alignment=TA_CENTER,
    )
    styles["CoverBrand"] = ParagraphStyle(
        "CoverBrand",
        parent=ss["Normal"],
        fontName=FONT_B,
        fontSize=12,
        leading=15,
        textColor=BRAND_GREEN,
        alignment=TA_CENTER,
    )
    styles["NoticeTitle"] = ParagraphStyle(
        "NoticeTitle",
        parent=styles["Body"],
        fontName=FONT_B,
        fontSize=11,
        leading=14,
        textColor=HEADER_TEXT,
        spaceAfter=4,
    )
    styles["NoticeBody"] = ParagraphStyle(
        "NoticeBody",
        parent=styles["Body"],
        fontSize=10.5,
        leading=14,
        textColor=BODY_TEXT,
        spaceAfter=0,
    )
    styles["NoticeBodyCenter"] = ParagraphStyle(
        "NoticeBodyCenter",
        parent=styles["NoticeBody"],
        alignment=TA_CENTER,
    )
    styles["Cell"] = ParagraphStyle(
        "Cell",
        parent=styles["Body"],
        fontName=FONT,
        fontSize=9,
        leading=11.5,
        spaceAfter=0,
    )
    styles["CellBold"] = ParagraphStyle(
        "CellBold",
        parent=styles["Cell"],
        fontName=FONT_B,
    )
    styles["CellHead"] = ParagraphStyle(
        "CellHead",
        parent=styles["Cell"],
        fontName=FONT_B,
        textColor=BRAND_GREEN,
    )
    return styles


# ---------------------------------------------------------------------------
# Callout / notice box
# ---------------------------------------------------------------------------
def notice_box(title: Optional[str], body_flows) -> Table:
    """A single-cell bordered box with an amber border and pale amber fill,
    matching the v0.1 "DEMO PURPOSES ONLY" / "Important Notice" style."""
    inner = []
    if title:
        inner.append(Paragraph(title, STYLES["NoticeTitle"]))
        inner.append(Spacer(1, 2))
    for f in body_flows:
        inner.append(f)
    t = Table([[inner]], colWidths=[CONTENT_W])
    t.setStyle(TableStyle([
        ("BOX",         (0, 0), (-1, -1), 1.5, NOTICE_BORDER),
        ("BACKGROUND",  (0, 0), (-1, -1), NOTICE_BG),
        ("LEFTPADDING", (0, 0), (-1, -1), 12),
        ("RIGHTPADDING",(0, 0), (-1, -1), 12),
        ("TOPPADDING",  (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 10),
        ("VALIGN",      (0, 0), (-1, -1), "TOP"),
    ]))
    return t


# ---------------------------------------------------------------------------
# Data table (multi-cell). First row is treated as a header.
# ---------------------------------------------------------------------------
def data_table(rows: List[List[str]]) -> Table:
    n_cols = max(len(r) for r in rows)
    # Even column widths unless first col is a short "label" pattern
    # Heuristic: if col 0 is short-ish, give it 35% and rest split
    first_col_lens = [len(r[0] or "") for r in rows]
    avg_first = sum(first_col_lens) / len(first_col_lens) if first_col_lens else 0
    if n_cols == 2:
        col_widths = [CONTENT_W * 0.30, CONTENT_W * 0.70] if avg_first < 45 else [CONTENT_W * 0.42, CONTENT_W * 0.58]
    elif n_cols == 3:
        col_widths = [CONTENT_W * 0.22, CONTENT_W * 0.30, CONTENT_W * 0.48]
    else:
        col_widths = [CONTENT_W / n_cols] * n_cols

    body_rows = []
    for i, r in enumerate(rows):
        cells = []
        for j in range(n_cols):
            text = r[j] if j < len(r) else ""
            style = STYLES["CellHead"] if i == 0 else STYLES["Cell"]
            cells.append(Paragraph(text or "", style))
        body_rows.append(cells)

    t = Table(body_rows, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (-1, 0), CELL_HEAD_FILL),
        ("GRID",         (0, 0), (-1, -1), 0.5, RULE_GREY),
        ("BOX",          (0, 0), (-1, -1), 0.75, RULE_GREY),
        ("VALIGN",       (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING",  (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING",   (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
    ]))
    return t


# ---------------------------------------------------------------------------
# docx -> flowable extraction
# ---------------------------------------------------------------------------
def para_kind(p) -> str:
    """Classify a python-docx paragraph by its style name."""
    style = (p.style.name if p.style else "").strip()
    if style == "Heading 1":
        return "h1"
    if style == "Heading 2":
        return "h2"
    if style == "Heading 3":
        return "h3"
    if style == "List Paragraph":
        return "bullet"
    return "body"


def escape_xml(s: str) -> str:
    return (
        s.replace("&", "&amp;")
         .replace("<", "&lt;")
         .replace(">", "&gt;")
    )


def para_html(p) -> str:
    """Render a docx paragraph's runs as ReportLab mini-HTML,
    preserving bold and italic."""
    parts = []
    for run in p.runs:
        text = escape_xml(run.text or "")
        if not text:
            continue
        if run.bold and run.italic:
            parts.append(f"<b><i>{text}</i></b>")
        elif run.bold:
            parts.append(f"<b>{text}</b>")
        elif run.italic:
            parts.append(f"<i>{text}</i>")
        else:
            parts.append(text)
    return "".join(parts) or escape_xml(p.text or "")


def cell_text(cell) -> str:
    """Extract cell text preserving line breaks between paragraphs."""
    return "\n".join(p.text for p in cell.paragraphs).strip()


def table_to_rows(table) -> List[List[str]]:
    return [[cell_text(c) for c in row.cells] for row in table.rows]


def is_callout(table) -> bool:
    """Single-cell tables are used as callout / notice boxes."""
    return len(table.rows) == 1 and len(table.rows[0].cells) == 1


def callout_body_flows(table, styles) -> (Optional[str], list):
    """Extract title (first paragraph if bold) and body flowables from a
    single-cell callout table."""
    cell = table.rows[0].cells[0]
    paras = [p for p in cell.paragraphs]
    if not paras:
        return None, []

    title = None
    body_paras = paras
    first = paras[0]
    # Heuristic: if the whole first paragraph is bold, treat it as the title
    if first.runs and all((r.bold for r in first.runs if r.text.strip())):
        title = first.text.strip()
        body_paras = paras[1:]

    flows = []
    for i, p in enumerate(body_paras):
        html = para_html(p)
        if not html.strip():
            continue
        st = styles["NoticeBody"]
        flows.append(Paragraph(html, st))
        if i < len(body_paras) - 1:
            flows.append(Spacer(1, 3))
    return title, flows


def iter_body_children(doc: Document):
    """Yield ('para', paragraph_obj) and ('table', table_obj) in document
    order."""
    paras = list(doc.paragraphs)
    tables = list(doc.tables)
    p_idx = t_idx = 0
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield "para", paras[p_idx]
            p_idx += 1
        elif tag == "tbl":
            yield "table", tables[t_idx]
            t_idx += 1


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------
def build_cover(styles, version: str, date: str, logo_path: Optional[Path]):
    flows = []
    # Logo (small, centered above title) -- optional
    if logo_path and logo_path.exists():
        try:
            img = Image(str(logo_path), width=1.4 * inch, height=1.4 * inch)
            img.hAlign = "CENTER"
            flows.append(Spacer(1, 0.6 * inch))
            flows.append(img)
            flows.append(Spacer(1, 0.25 * inch))
        except Exception:
            flows.append(Spacer(1, 0.9 * inch))
    else:
        flows.append(Spacer(1, 1.4 * inch))

    flows.append(Paragraph("Financed Emissions Dashboard", styles["CoverTitle"]))
    flows.append(Spacer(1, 6))
    flows.append(Paragraph("User's Manual", styles["CoverSub"]))
    flows.append(Spacer(1, 30))
    flows.append(Paragraph(
        "Reference demonstration for commercial banks in Nepal",
        styles["CoverTag"],
    ))
    flows.append(Spacer(1, 8))
    flows.append(Paragraph(
        f"Demonstration Build &nbsp;|&nbsp; v{version} (Demo Build)",
        styles["CoverMeta"],
    ))
    flows.append(Spacer(1, 4))
    flows.append(Paragraph(date, styles["CoverMeta"]))
    flows.append(Spacer(1, 0.55 * inch))
    flows.append(Paragraph("Published by Jana Earth Data", styles["CoverBrand"]))
    flows.append(Spacer(1, 0.55 * inch))
    flows.append(notice_box(
        "DEMO PURPOSES ONLY",
        [Paragraph(
            "This document accompanies a non-production demonstration build of the "
            "Financed Emissions Dashboard, made available as a reference to "
            "commercial banks in Nepal. The example bank shown in the dashboard is "
            "hypothetical and is used only to make the screens concrete.",
            styles["NoticeBodyCenter"],
        )],
    ))
    flows.append(PageBreak())
    return flows


def build_body_flows(doc: Document, styles) -> list:
    """Convert every docx body element (except the cover block we render
    ourselves) into a list of ReportLab flowables."""
    flows = []
    started = False  # skip everything until the first Heading 1 ("Important Notice")

    for kind, node in iter_body_children(doc):
        if not started:
            if kind == "para" and para_kind(node) == "h1":
                started = True
            else:
                continue

        if kind == "para":
            k = para_kind(node)
            html = para_html(node)
            text = node.text or ""
            if k == "h1":
                if flows:
                    flows.append(PageBreak())
                flows.append(Paragraph(html or escape_xml(text), styles["H1"]))
                # A green hairline under Part headings
                flows.append(Spacer(1, 2))
            elif k == "h2":
                flows.append(Paragraph(html or escape_xml(text), styles["H2"]))
            elif k == "h3":
                flows.append(Paragraph(html or escape_xml(text), styles["H3"]))
            elif k == "bullet":
                if not text.strip():
                    continue
                flows.append(Paragraph(html, styles["Bullet"], bulletText="•"))
            else:
                if not text.strip():
                    flows.append(Spacer(1, 4))
                    continue
                flows.append(Paragraph(html, styles["Body"]))
        else:  # table
            if is_callout(node):
                title, body = callout_body_flows(node, styles)
                flows.append(Spacer(1, 4))
                flows.append(notice_box(title, body))
                flows.append(Spacer(1, 6))
            else:
                rows = table_to_rows(node)
                flows.append(Spacer(1, 2))
                flows.append(data_table(rows))
                flows.append(Spacer(1, 6))
    return flows


# ---------------------------------------------------------------------------
# Page templates: cover (no chrome) + body (running header + footer)
# ---------------------------------------------------------------------------
class ManualDocTemplate(BaseDocTemplate):
    def __init__(self, filename, *, version: str, date: str, **kwargs):
        super().__init__(
            filename,
            pagesize=LETTER,
            leftMargin=MARGIN_L,
            rightMargin=MARGIN_R,
            topMargin=MARGIN_T,
            bottomMargin=MARGIN_B,
            **kwargs,
        )
        self.version = version
        self.date = date
        self._page_num = 0

        cover_frame = Frame(
            MARGIN_L, MARGIN_B, CONTENT_W, PAGE_H - MARGIN_T - MARGIN_B,
            id="cover", leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0,
        )
        body_frame = Frame(
            MARGIN_L, MARGIN_B + 0.35 * inch, CONTENT_W,
            PAGE_H - MARGIN_T - MARGIN_B - 0.7 * inch,
            id="body", leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0,
        )
        self.addPageTemplates([
            PageTemplate(id="cover", frames=[cover_frame], onPage=self._draw_cover_chrome),
            PageTemplate(id="body",  frames=[body_frame],  onPage=self._draw_body_chrome),
        ])

    # --- cover chrome: just the footer tag; no running header ------------
    def _draw_cover_chrome(self, canvas, doc):
        canvas.saveState()
        canvas.setFont(FONT, 8.5)
        canvas.setFillColor(FOOTER_GREY)
        footer = (
            f"DEMO PURPOSES ONLY  |  Synthesized portfolio against real public "
            f"reference data"
        )
        canvas.drawString(MARGIN_L, MARGIN_B - 0.3 * inch, footer)
        canvas.drawRightString(
            PAGE_W - MARGIN_R, MARGIN_B - 0.3 * inch,
            f"Page 1  |  {self.date}",
        )
        canvas.restoreState()

    # --- body chrome: two-line header + footer with page num -------------
    def _draw_body_chrome(self, canvas, doc):
        canvas.saveState()
        # Header
        canvas.setFont(FONT, 8.5)
        canvas.setFillColor(HEADER_TEXT)
        header_left  = "Financed Emissions Dashboard  |  Demonstration Build"
        header_right = f"User's Manual  |  v{self.version} (Demo Build)"
        y_head = PAGE_H - MARGIN_T + 0.35 * inch
        canvas.drawString(MARGIN_L, y_head, header_left)
        canvas.drawRightString(PAGE_W - MARGIN_R, y_head, header_right)
        # Hairline under header
        canvas.setStrokeColor(BRAND_GREEN)
        canvas.setLineWidth(0.6)
        canvas.line(MARGIN_L, y_head - 4, PAGE_W - MARGIN_R, y_head - 4)

        # Footer
        canvas.setFont(FONT, 8.5)
        canvas.setFillColor(FOOTER_GREY)
        footer_left = (
            "DEMO PURPOSES ONLY  |  Synthesized portfolio against real public "
            "reference data"
        )
        y_foot = MARGIN_B - 0.3 * inch
        canvas.drawString(MARGIN_L, y_foot, footer_left)
        canvas.drawRightString(
            PAGE_W - MARGIN_R, y_foot,
            f"Page {doc.page}  |  {self.date}",
        )
        # Hairline above footer
        canvas.setStrokeColor(RULE_GREY)
        canvas.setLineWidth(0.4)
        canvas.line(MARGIN_L, y_foot + 12, PAGE_W - MARGIN_R, y_foot + 12)
        canvas.restoreState()

    def afterFlowable(self, flowable):
        # No-op — hook available for future TOC generation
        return


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def parse_args():
    repo_root = Path(__file__).resolve().parent.parent
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--version", default="0.2",
                    help="Manual version, e.g. 0.2 (default: 0.2)")
    ap.add_argument("--date", default="August 2026",
                    help="Publication date shown on cover and in footer")
    ap.add_argument("--docx", type=Path,
                    default=repo_root / "docs"
                    / "Jana_Financed_Emissions_Dashboard_User_Manual_Demo_v0.2_DRAFT.docx",
                    help="Source docx path")
    ap.add_argument("--out",  type=Path,
                    default=repo_root / "docs"
                    / "Jana_Financed_Emissions_Dashboard_User_Manual_Demo_v0.2.pdf",
                    help="Output PDF path")
    ap.add_argument("--logo", type=Path,
                    default=repo_root / "public" / "green_logo.png",
                    help="Cover logo image (optional)")
    return ap.parse_args()


def main():
    global STYLES
    args = parse_args()

    if not args.docx.exists():
        sys.exit(f"error: docx source not found: {args.docx}")

    doc = Document(str(args.docx))
    STYLES = build_styles()

    flows = []
    flows.extend(build_cover(STYLES, args.version, args.date,
                             args.logo if args.logo and args.logo.exists() else None))
    flows.extend(build_body_flows(doc, STYLES))

    args.out.parent.mkdir(parents=True, exist_ok=True)
    template = ManualDocTemplate(
        str(args.out),
        version=args.version,
        date=args.date,
        title=f"Financed Emissions Dashboard User's Manual",
        subject="Reference demonstration for commercial banks in Nepal",
        author="Jana Earth Data",
    )

    # First page uses the cover template; subsequent pages use body
    # (achieved by having the cover flows end in a PageBreak, then setting
    # the next template.)
    from reportlab.platypus.doctemplate import NextPageTemplate

    # Insert the NextPageTemplate("body") right before the final PageBreak
    # of the cover block so page 2 flips to the body template.
    for i, f in enumerate(flows):
        if isinstance(f, PageBreak):
            flows.insert(i, NextPageTemplate("body"))
            break

    template.build(flows)

    size_kb = args.out.stat().st_size / 1024
    print(f"Wrote {args.out}  ({size_kb:.1f} KB)")


if __name__ == "__main__":
    STYLES = {}  # populated in main()
    main()
